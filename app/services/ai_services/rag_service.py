from app.utils.ai_client import ai_client
from app.utils.milvus_client import milvus_client
from app.config.settings import settings
from app.services.golden_qa_service import golden_qa_service
from app.services.engineer_level_service import engineer_level_service
from app.config.constants import EngineerLevel
import os
import re
import json
import logging
import shutil
import subprocess
import tempfile
import requests
import numpy as np

logger = logging.getLogger(__name__)

# 扫描版 PDF 逐页 OCR 上限（避免超大 PDF 消耗过多 API）
PDF_OCR_MAX_PAGES = 30


class SiliconFlowEmbedder:
    """硅基流动 OpenAI 兼容 embedding API"""

    def __init__(self, model: str | None = None):
        self.model = model or settings.SILICON_FLOW_EMBEDDING_MODEL
        self.api_key = settings.SILICON_FLOW_API_KEY
        self.url = settings.SILICON_FLOW_EMBEDDING_URL.rstrip("/")
        self.dim: int | None = settings.EMBEDDING_DIM or None

    def encode(self, texts, **kwargs):
        if not self.api_key:
            raise ValueError("未配置 SILICON_FLOW_API_KEY，无法使用硅基流动向量化")

        single = isinstance(texts, str)
        inputs = [texts] if single else texts
        response = requests.post(
            self.url,
            headers={
                "Authorization": f"Bearer {self.api_key}",
                "Content-Type": "application/json",
            },
            json={"model": self.model, "input": inputs},
            timeout=120,
        )
        if response.status_code >= 400:
            detail = response.text[:300]
            try:
                body = response.json()
                detail = body.get("message") or body.get("error") or detail
            except Exception:
                pass
            raise RuntimeError(
                f"硅基流动 embedding 失败 (HTTP {response.status_code}): {detail}"
            )

        data = response.json()
        items = sorted(data.get("data", []), key=lambda item: item.get("index", 0))
        embeddings = [item["embedding"] for item in items]
        if not embeddings:
            raise RuntimeError("硅基流动 embedding 返回为空")

        if self.dim is None:
            self.dim = len(embeddings[0])

        result = np.array(embeddings, dtype=np.float32)
        return result[0] if single else result


class RAGService:
    _instance = None

    def __new__(cls):
        if cls._instance is None:
            cls._instance = super().__new__(cls)
            cls._instance._initialized = False
        return cls._instance

    def __init__(self):
        if self._initialized:
            return
        self._initialized = True
        self._embedding_model = None
        self._text_splitter = None

    def _load_embedding(self):
        if self._embedding_model is None:
            logger.info(
                "Using SiliconFlow embedding model: %s (dim=%s)",
                settings.SILICON_FLOW_EMBEDDING_MODEL,
                settings.EMBEDDING_DIM,
            )
            self._embedding_model = SiliconFlowEmbedder()

    def _embedding_dim(self) -> int:
        return settings.EMBEDDING_DIM

    def _load_splitter(self):
        if self._text_splitter is None:
            from langchain_text_splitters import RecursiveCharacterTextSplitter
            self._text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)

    def _ensure_collections(self):
        dim = self._embedding_dim()
        self.knowledge_collection = milvus_client.ensure_knowledge_collection(dim=dim)
        self.cases_collection = milvus_client.ensure_cases_collection(dim=dim)

    async def ingest_manual(self, file_path: str, source: str = "",
                             doc_level: str = EngineerLevel.JUNIOR,
                             equipment_id: str = "") -> list:
        self._load_splitter()
        self._load_embedding()
        self._ensure_collections()

        text = await self._extract_text(file_path)
        if not text.strip():
            raise ValueError(f"未能从文件中提取到文本内容: {os.path.basename(file_path)}")
        chunks = self._text_splitter.split_text(text)

        count = 0
        for chunk in chunks:
            embedding = self._embedding_model.encode(chunk).tolist()
            self.knowledge_collection.insert([{
                "content": chunk,
                "source": source or os.path.basename(file_path),
                "doc_level": doc_level,
                "equipment_id": equipment_id,
                "embedding": embedding,
            }])
            count += 1
        self.knowledge_collection.flush()
        return list(range(count))

    async def query_knowledge(self, question: str, top_k: int = 3,
                               user_level: str = EngineerLevel.JUNIOR,
                               equipment_id: str = "") -> dict:
        self._load_embedding()
        self._ensure_collections()

        query_embedding = self._embedding_model.encode(question).tolist()
        self.knowledge_collection.load()

        # 黄金问答优先注入
        golden_items = await golden_qa_service.search(question, top_k=3, scene="knowledge")

        # 按工程师等级过滤可见文档级别
        visible_levels = engineer_level_service.get_visible_doc_levels(user_level)

        search_params = {"metric_type": "COSINE", "params": {"nprobe": 16}}
        results = self.knowledge_collection.search(
            data=[query_embedding],
            anns_field="embedding",
            param=search_params,
            limit=top_k + 5,
            output_fields=["content", "source", "doc_level", "equipment_id"],
        )

        docs = []
        for hits in results:
            for hit in hits:
                doc_level = hit.entity.get("doc_level", EngineerLevel.JUNIOR)
                if doc_level not in visible_levels:
                    continue
                if equipment_id and hit.entity.get("equipment_id", "") != equipment_id:
                    continue
                docs.append({
                    "content": hit.entity.get("content"),
                    "source": hit.entity.get("source"),
                    "distance": hit.distance,
                })
                if len(docs) >= top_k:
                    break

        # 构建上下文：优先注入黄金问答
        context_parts = []
        for item in golden_items:
            context_parts.append(f"{item['问题']}\n{item['标准答案']}")
        for d in docs:
            context_parts.append(d["content"])
        context = "\n\n".join(context_parts)

        prompt = f"""你是设备维修专家助手。请根据以下设备资料和标准答案回答问题。

【设备资料/标准答案】
{context}

【用户问题】
{question}

请给出详细排查步骤，并注明信息来源。"""

        answer = await ai_client.chat(prompt)

        return {
            "answer": answer,
            "sources": [
                {
                    "source": d["source"],
                    "content": d["content"],
                    "distance": d["distance"],
                }
                for d in docs
            ],
            "golden_qa": golden_items,
        }

    async def query_knowledge_stream(self, question: str, top_k: int = 3,
                                      user_level: str = EngineerLevel.JUNIOR,
                                      equipment_id: str = ""):
        """流式版知识检索"""
        result = await self.query_knowledge(question, top_k=top_k,
                                             user_level=user_level,
                                             equipment_id=equipment_id)
        sources = result.get("sources", [])
        async for event in ai_client.chat_stream_with_sources(question, sources=sources):
            yield event

    async def search_fault_cases(self, query: str, top_k: int = 5) -> dict:
        self._load_embedding()
        self._ensure_collections()

        query_embedding = self._embedding_model.encode(query).tolist()
        self.cases_collection.load()

        search_params = {"metric_type": "L2", "params": {"nprobe": 16}}
        results = self.cases_collection.search(
            data=[query_embedding],
            anns_field="embedding",
            param=search_params,
            limit=top_k,
            output_fields=["content", "fault_type", "resolution"],
        )

        cases = []
        for hits in results:
            for hit in hits:
                cases.append({
                    "content": hit.entity.get("content"),
                    "fault_type": hit.entity.get("fault_type"),
                    "resolution": hit.entity.get("resolution"),
                    "distance": hit.distance,
                })

        return {"cases": cases}

    async def ingest_fault_case(self, fault_desc: str, fault_type: str, resolution: str):
        self._load_embedding()
        self._ensure_collections()

        embedding = self._embedding_model.encode(fault_desc).tolist()
        self.cases_collection.insert([{
            "content": fault_desc,
            "fault_type": fault_type,
            "resolution": resolution,
            "embedding": embedding,
        }])
        self.cases_collection.flush()

    async def ingest_text(self, text: str, metadata: dict = None):
        """直接录入文本到知识库（如维修报告自动入库）"""
        self._load_embedding()
        self._ensure_collections()

        metadata = metadata or {}
        source = metadata.get("source", "manual")
        doc_level = metadata.get("doc_level", EngineerLevel.JUNIOR)
        equipment_id = metadata.get("equipment_id", "")
        self._load_splitter()
        chunks = self._text_splitter.split_text(text)
        for chunk in chunks:
            embedding = self._embedding_model.encode(chunk).tolist()
            self.knowledge_collection.insert([{
                "content": chunk,
                "source": source,
                "doc_level": doc_level,
                "equipment_id": equipment_id,
                "embedding": embedding,
            }])
        self.knowledge_collection.flush()

    async def _extract_text(self, file_path: str) -> str:
        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".pdf":
            return await self._extract_pdf(file_path)
        if ext in (".docx", ".doc"):
            return self._extract_office(file_path, ext)
        if ext in (".png", ".jpg", ".jpeg", ".bmp", ".webp"):
            return await self._extract_image_text(file_path)
        return self._extract_plain_text(file_path)

    def _detect_office_format(self, file_path: str) -> str:
        """按文件头识别真实格式，避免 .doc 改后缀为 .docx"""
        with open(file_path, "rb") as f:
            header = f.read(8)
        if header[:2] == b"PK":
            return "docx"
        if len(header) >= 8 and header[:8] == b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1":
            return "doc"
        return "unknown"

    def _extract_office(self, file_path: str, ext: str) -> str:
        fmt = self._detect_office_format(file_path)
        if fmt == "docx":
            return self._extract_docx(file_path)
        if fmt == "doc":
            return self._extract_doc(file_path)
        # 无法识别文件头时按扩展名尝试，docx 失败则回退 doc 解析链
        if ext == ".docx":
            try:
                return self._extract_docx(file_path)
            except Exception as exc:
                logger.info("按 docx 打开失败，尝试 doc 解析链: %s", exc)
                return self._extract_doc(file_path)
        return self._extract_doc(file_path)

    async def _extract_pdf(self, file_path: str) -> str:
        text = self._extract_pdf_text_layer(file_path)
        if text.strip():
            return text
        logger.info("PDF 无文本层或提取为空，改用硅基流动逐页 OCR: %s", os.path.basename(file_path))
        return await self._extract_pdf_via_vision(file_path)

    def _extract_pdf_text_layer(self, file_path: str) -> str:
        """从 PDF 文本层提取（pdfplumber + PyMuPDF）"""
        text_parts: list[str] = []

        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
        except ImportError:
            raise ImportError("pdfplumber 未安装，请执行: pip install pdfplumber")
        except Exception as exc:
            logger.debug("pdfplumber 解析失败: %s", exc)

        text = "\n".join(text_parts)
        if text.strip():
            return text

        try:
            import fitz
            doc = fitz.open(file_path)
            try:
                for page in doc:
                    page_text = page.get_text()
                    if page_text and page_text.strip():
                        text_parts.append(page_text)
            finally:
                doc.close()
        except ImportError:
            logger.debug("pymupdf 未安装，跳过 PyMuPDF 文本提取")
        except Exception as exc:
            logger.debug("PyMuPDF 文本提取失败: %s", exc)

        return "\n".join(text_parts)

    async def _extract_pdf_via_vision(self, file_path: str) -> str:
        """扫描版 PDF：逐页渲染为图片，硅基流动视觉模型 OCR"""
        try:
            import fitz
        except ImportError:
            raise ImportError("pymupdf 未安装，无法 OCR 扫描版 PDF，请执行: pip install pymupdf")

        parts: list[str] = []
        doc = fitz.open(file_path)
        try:
            total = doc.page_count
            limit = min(total, PDF_OCR_MAX_PAGES)
            if total > PDF_OCR_MAX_PAGES:
                logger.warning(
                    "PDF 共 %s 页，仅 OCR 前 %s 页: %s",
                    total,
                    PDF_OCR_MAX_PAGES,
                    os.path.basename(file_path),
                )
            for page_index in range(limit):
                page = doc.load_page(page_index)
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                image_bytes = pix.tobytes("png")
                page_text = await ai_client.extract_document_text_from_image(
                    image_bytes,
                    mime_type="image/png",
                )
                if page_text.strip():
                    parts.append(page_text.strip())
        finally:
            doc.close()

        text = "\n\n".join(parts)
        if not text.strip():
            raise ValueError(
                f"扫描版 PDF OCR 未识别到文字: {os.path.basename(file_path)}，"
                "请确认 PDF 清晰可读"
            )
        return text

    def _extract_docx(self, file_path: str) -> str:
        try:
            from docx import Document
            from docx.opc.exceptions import PackageNotFoundError
        except ImportError:
            raise ImportError("python-docx 未安装，请执行: pip install python-docx")
        try:
            doc = Document(file_path)
        except PackageNotFoundError as exc:
            raise ValueError(
                "文件扩展名为 .docx，但不是有效的 Word 文档（可能是旧版 .doc 改后缀），"
                "系统将尝试其他解析方式"
            ) from exc
        parts: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append("\t".join(cells))
        text = "\n".join(parts)
        if not text.strip():
            raise ValueError("Word 文档未提取到文本内容")
        return text

    def _extract_doc(self, file_path: str) -> str:
        """解析旧版 .doc：Word COM → LibreOffice → OLE 流 → antiword"""
        abs_path = os.path.abspath(file_path)
        parsers = (
            self._extract_doc_via_word,
            self._extract_doc_via_libreoffice,
            self._extract_doc_via_olefile,
            self._extract_doc_via_antiword,
        )
        for parser in parsers:
            try:
                text = parser(abs_path)
                if text and text.strip():
                    return text.strip()
            except Exception as exc:
                logger.debug(".doc 解析 %s 失败: %s", parser.__name__, exc)
        raise ValueError(
            "无法解析 .doc 文件，请安装 Microsoft Word 或 LibreOffice，或另存为 .docx 后重新上传"
        )

    def _extract_doc_via_word(self, abs_path: str) -> str:
        import win32com.client

        word = None
        doc = None
        try:
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            doc = word.Documents.Open(abs_path, ReadOnly=True, AddToRecentFiles=False)
            return (doc.Content.Text or "").replace("\r", "\n")
        finally:
            if doc is not None:
                doc.Close(False)
            if word is not None:
                word.Quit()

    def _find_soffice(self) -> str | None:
        found = shutil.which("soffice")
        if found:
            return found
        for candidate in (
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ):
            if os.path.isfile(candidate):
                return candidate
        return None

    def _extract_doc_via_libreoffice(self, abs_path: str) -> str:
        soffice = self._find_soffice()
        if not soffice:
            raise FileNotFoundError("未找到 LibreOffice")

        out_dir = tempfile.mkdtemp()
        try:
            result = subprocess.run(
                [soffice, "--headless", "--convert-to", "txt:Text", "--outdir", out_dir, abs_path],
                capture_output=True,
                timeout=120,
                check=False,
            )
            if result.returncode != 0:
                stderr = result.stderr.decode("utf-8", errors="ignore")
                raise RuntimeError(stderr or "LibreOffice 转换失败")

            for name in os.listdir(out_dir):
                if name.lower().endswith(".txt"):
                    txt_path = os.path.join(out_dir, name)
                    with open(txt_path, "r", encoding="utf-8", errors="ignore") as f:
                        return f.read()
            raise ValueError("LibreOffice 未生成 txt 文件")
        finally:
            shutil.rmtree(out_dir, ignore_errors=True)

    def _extract_doc_via_olefile(self, abs_path: str) -> str:
        import olefile

        if not olefile.isOleFile(abs_path):
            raise ValueError("不是有效的 OLE .doc 文件")

        chunks: list[str] = []
        ole = olefile.OleFileIO(abs_path)
        try:
            if not ole.exists("WordDocument"):
                raise ValueError("缺少 WordDocument 流")
            data = ole.openstream("WordDocument").read()
            for match in re.finditer(rb"[\x20-\x7e\r\n\t]{4,}", data):
                chunks.append(match.group().decode("ascii"))
            u16_text = data.decode("utf-16-le", errors="ignore")
            u16_text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", " ", u16_text)
            for match in re.finditer(
                r"[\u4e00-\u9fff\u3000-\u303f\uff00-\uffef\w\s，。；：、（）\"\"''？！·\-—]{4,}",
                u16_text,
            ):
                chunks.append(match.group())
        finally:
            ole.close()

        text = "\n".join(dict.fromkeys(part.strip() for part in chunks if part.strip()))
        if len(text) < 10:
            raise ValueError("OLE 流未提取到足够文本")
        return text

    def _extract_doc_via_antiword(self, abs_path: str) -> str:
        result = subprocess.run(
            ["antiword", abs_path],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="ignore",
            timeout=60,
            check=False,
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout
        raise RuntimeError("antiword 解析失败")

    def _extract_plain_text(self, file_path: str) -> str:
        for encoding in ("utf-8", "gbk", "gb2312", "latin-1"):
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        raise ValueError(f"无法读取文本文件编码: {os.path.basename(file_path)}")

    async def _extract_image_text(self, file_path: str) -> str:
        """图片 OCR：统一使用硅基流动视觉模型"""
        mime_type = ai_client.get_image_mime_type(file_path)
        with open(file_path, "rb") as f:
            image_bytes = f.read()
        text = await ai_client.extract_document_text_from_image(image_bytes, mime_type=mime_type)
        if not text.strip():
            raise ValueError("图片 OCR 未识别到文字，请确认图片清晰且包含可读文本")
        return text

    # ========== 多模态 PDF 处理（图文混合切片） ==========

    async def ingest_pdf_multimodal(self, file_path: str, source: str = "") -> dict:
        """
        多模态 PDF 入库流水线：
        PDF → 版面分析 → 图文混合切片 → 向量化 → 存入 Milvus 多模态集合
        返回入库统计信息。
        """
        from app.utils.pdf_layout_parser import pdf_layout_parser
        from app.utils.mixed_chunker import create_mixed_chunks

        self._load_embedding()
        self._ensure_multimodal_collection()

        # 1. 版面分析
        blocks = pdf_layout_parser.parse(file_path)
        if not blocks:
            raise ValueError(f"未能从 PDF 中解析到任何内容: {os.path.basename(file_path)}")

        # 2. 图文混合切片
        chunks = create_mixed_chunks(blocks, source=source)
        if not chunks:
            raise ValueError("图文混合切片后未生成任何 Chunk")

        # 3. 向量化 + 入库
        inserted = 0
        for chunk in chunks:
            try:
                embedding = self._embedding_model.encode(chunk.text).tolist()
                self.multimodal_collection.insert([{
                    "content": chunk.text,
                    "source": source,
                    "chunk_type": chunk.chunk_type,
                    "page_num": chunk.page_num,
                    "bbox": chunk.bbox,
                    "image_path": chunk.image_path,
                    "image_hash": chunk.image_hash,
                    "embedding": embedding,
                }])
                inserted += 1
            except Exception as exc:
                logger.warning("Chunk 入库失败 (type=%s, page=%s): %s",
                               chunk.chunk_type, chunk.page_num, exc)
                continue

        self.multimodal_collection.flush()

        # 4. 统计
        stats = {
            "total_chunks": len(chunks),
            "inserted": inserted,
            "by_type": {},
            "with_images": 0,
        }
        for chunk in chunks:
            t = chunk.chunk_type
            stats["by_type"][t] = stats["by_type"].get(t, 0) + 1
            if chunk.image_path:
                stats["with_images"] += 1

        return stats

    @staticmethod
    async def decide_upload_strategy(
        filename: str,
        file_size: int,
        mime_type: str = "",
        pdf_page_count: int = 0,
        pdf_has_images: bool = False,
        text_preview: str = "",
    ) -> str:
        """
        由大模型判断文件该用哪种处理策略。

        返回 "simple_text" 或 "multimodal_layout"。
        """
        ext = os.path.splitext(filename)[1].lower()

        # ── 规则兜底：非 PDF 文件不需要版面分析 ──
        if ext not in (".pdf",):
            return "simple_text"

        # ── PDF 但无嵌入图片 → simple_text ──
        if not pdf_has_images:
            return "simple_text"

        # ── PDF 有图片，让大模型根据内容判断 ──
        preview_text = text_preview[:500] if text_preview else "（无预览）"

        prompt = (
            "你是一个文档处理策略专家。请根据以下文件信息，判断最适合的处理策略。\n\n"
            f"【文件信息】\n"
            f"- 文件名: {filename}\n"
            f"- 大小: {file_size / 1024:.1f} KB\n"
            f"- 类型: PDF\n"
            f"- 页数: {pdf_page_count}\n"
            f"- 包含嵌入图片: 是\n\n"
            f"【内容预览（前500字）】\n{preview_text}\n\n"
            f"【可选策略】\n"
            f"1. simple_text —— 简单文本提取 + RecursiveCharacterTextSplitter 切片\n"
            f"   适用场景：纯文本 PDF，内容以段落文字为主，图片只是装饰性元素\n\n"
            f"2. multimodal_layout —— 版面分析 + 图文混合切片 + 图片 OCR\n"
            f"   适用场景：包含技术图表、设备照片、示意图、流程图等技术文档 PDF，\n"
            f"   图片承载了重要信息，需要图文绑定以保留语义\n\n"
            "请只输出策略名称（simple_text 或 multimodal_layout），不要输出任何其他内容。"
        )
        try:
            result = await ai_client.chat(prompt, temperature=0.1)
            result = result.strip().lower()
            if "multimodal_layout" in result:
                return "multimodal_layout"
            return "simple_text"
        except Exception as exc:
            logger.warning("LLM 策略判断失败，默认使用 simple_text: %s", exc)
            return "simple_text"

    async def query_multimodal(
        self,
        question: str,
        top_k: int = 5,
        return_images: bool = True,
    ) -> dict:
        """
        多模态知识库检索。

        返回：
        - answer: LLM 基于检索上下文生成的回答
        - sources: 匹配的 Chunk 列表（含文本、来源、图片路径等）
        - images: 图片的 MinIO 预签名 URL 列表（可用于多模态 LLM 输入）
        """
        self._load_embedding()
        self._ensure_multimodal_collection()

        query_embedding = self._embedding_model.encode(question).tolist()
        self.multimodal_collection.load()

        search_params = {"metric_type": "L2", "params": {"nprobe": 16}}
        results = self.multimodal_collection.search(
            data=[query_embedding],
            anns_field="embedding",
            param=search_params,
            limit=top_k,
            output_fields=[
                "content", "source", "chunk_type",
                "page_num", "bbox", "image_path", "image_hash",
            ],
        )

        chunks = []
        image_urls = []
        seen_hashes = set()

        for hits in results:
            for hit in hits:
                content = hit.entity.get("content", "")
                image_path = hit.entity.get("image_path", "")
                img_hash = hit.entity.get("image_hash", "")

                chunk_info = {
                    "content": content,
                    "source": hit.entity.get("source", ""),
                    "chunk_type": hit.entity.get("chunk_type", ""),
                    "page_num": hit.entity.get("page_num", -1),
                    "bbox": hit.entity.get("bbox", ""),
                    "image_path": image_path,
                    "distance": hit.distance,
                }
                chunks.append(chunk_info)

                # 收集去重的图片 URL
                if return_images and image_path and img_hash and img_hash not in seen_hashes:
                    seen_hashes.add(img_hash)
                    try:
                        url = minio_client.get_file_url(
                            image_path,
                            bucket=settings.KB_BUCKET_NAME,
                        )
                        image_urls.append({
                            "image_path": image_path,
                            "presigned_url": url,
                            "page_num": hit.entity.get("page_num", -1),
                        })
                    except Exception as exc:
                        logger.warning("获取图片 URL 失败: %s", exc)

        # 用检索到的上下文生成回答
        context = "\n\n".join([c["content"] for c in chunks])
        prompt = (
            "你是设备维修专家助手。请根据以下图文资料回答问题。\n\n"
            f"【资料】\n{context}\n\n"
            f"【问题】\n{question}\n\n"
            "请给出详细回答，并注明信息来源页码。"
        )
        answer = await ai_client.chat(prompt)

        return {
            "answer": answer,
            "sources": chunks,
            "images": image_urls,
        }

    def _ensure_multimodal_collection(self):
        """确保多模态集合已初始化"""
        dim = self._embedding_dim()
        self.multimodal_collection = milvus_client.ensure_multimodal_collection(dim=dim)


rag_service = RAGService()